"""
PBAS Form Generator - Creates PBAS-S2.docx with exact formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class PBASFormGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
    
    def setup_document(self):
        """Setup document margins and default styles"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def add_title(self):
        """Add main title with bold and underline"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = title.add_run('Self- Assessment-Cum-Performance Appraisal Forms\nAPI-PBAS Proforma')
        run.bold = True
        run.underline = WD_UNDERLINE.SINGLE
        run.font.size = Pt(14)
    
    def add_basic_info_section(self, data):
        """Add basic information section"""
        self.doc.add_paragraph()
        
        # Institute name
        p = self.doc.add_paragraph()
        p.add_run('Name of the lnstitute / College: ').bold = True
        p.add_run(data.get('institute', '.'*50))
        
        # Department
        p = self.doc.add_paragraph()
        p.add_run('Name of fie Department: ').bold = True
        p.add_run(data.get('department', '.'*50))
        
        # CAS Promotion
        p = self.doc.add_paragraph()
        p.add_run('Under CAS Promotion for Stage/ Level').bold = True
        
        # For section
        p = self.doc.add_paragraph()
        p.add_run('For').italic = True
        p.add_run('\nFaculty of').italic = True
        p.add_run(data.get('faculty', '.'*50))
        
        # References
        p = self.doc.add_paragraph()
        p.add_run('Reference: ').bold = False
        ref_text = """i) The Gazette of India: Extraordinary, Part lII Section4 dated 18th July, 2018
ii) Government of Maharashtra Misc. - 2018.CR 56/18/UNI1date 8th March, 2019
iii) Government of Maharashtra Misc.- 20l 8.CR56/l 8/UNI I date 10th May, 2019"""
        p.add_run(ref_text)
        
        # Academic Year
        p = self.doc.add_paragraph()
        p.add_run('ACADEMIC YEAR: ').bold = True
        p.add_run(data.get('academic_year', '-'*20))
    
    def add_part_a_header(self):
        """Add Part A header"""
        p = self.doc.add_paragraph()
        p.add_run('PART A: GENERAL INFORMATION AND ACADEMIC BACKGROUND').bold = True
    
    def add_general_info_fields(self, data):
        """Add numbered fields 1-8"""
        fields = [
            ('1. Name (in Block Letters)', data.get('name', '')),
            ('2. Department', data.get('department', '')),
            ('3. Current Designation & Academic Level', data.get('current_designation', '')),
            ('4. Date of last Promotion Current position and Academic', data.get('last_promotion', '')),
            ('5. Level of an applicant under CAS', data.get('cas_level', '')),
            ('6. The designation and grade pay applied for under CAS', data.get('applied_designation', '')),
            ('7. Date of eligibility for promotion', data.get('eligibility_date', '')),
        ]
        
        for field_name, value in fields:
            p = self.doc.add_paragraph()
            p.add_run(field_name).bold = False
            if value:
                p.add_run('\n' + value)
        
        # Address field
        p = self.doc.add_paragraph()
        p.add_run('8. Address (with Pin code)').bold = False
        p.add_run('\n' + data.get('address', ''))
        p.add_run('\n\nTelephone/ Mobile No.')
        p.add_run('\n' + data.get('mobile', ''))
        p.add_run('\n\nE-mail')
        p.add_run('\n' + data.get('email', ''))
    
    def add_academic_qualifications_table(self, qualifications):
        """Add academic qualifications table (Point 9)"""
        p = self.doc.add_paragraph()
        p.add_run('Academic Qualifications (from S.S.C. till Post-Graduation):').bold = True
        
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Examinations', 'Name of the Board/University', 'Year of Passing', 
                   'Percentage Of Marks Obtained', 'Division I Class I Grade', 'Subject']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for qual in qualifications:
            row_cells = table.add_row().cells
            row_cells[0].text = qual.get('examination', '')
            row_cells[1].text = qual.get('board', '')
            row_cells[2].text = qual.get('year', '')
            row_cells[3].text = qual.get('percentage', '')
            row_cells[4].text = qual.get('division', '')
            row_cells[5].text = qual.get('subject', '')
    
    def add_research_degrees_table(self, degrees):
        """Add research degrees table (Point 10)"""
        p = self.doc.add_paragraph()
        p.add_run('10. Research Degree(s):').bold = True
        
        table = self.doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Degrees', 'Title', 'Date of Award', 'Name of University']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Degree types
        degree_types = ['M. Phil.', 'Ph.D. ID. Phil.', 'D.Sc. / D .Litt./ Any other']
        for i, deg_type in enumerate(degree_types, 1):
            table.rows[i].cells[0].text = deg_type
            # Fill data if provided
            if i-1 < len(degrees):
                table.rows[i].cells[1].text = degrees[i-1].get('title', '')
                table.rows[i].cells[2].text = degrees[i-1].get('date', '')
                table.rows[i].cells[3].text = degrees[i-1].get('university', '')
    
    def add_prior_appointments_table(self, appointments):
        """Add prior appointments table (Point 11)"""
        p = self.doc.add_paragraph()
        p.add_run('11. Appointments held prior-joining this institution:(Please attach relevant certificates of service/experience)').bold = True
        
        table = self.doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        
        headers = ['Designation', 'Name of Employer', 'Essential Qualifications for the Post at the time of Apointment',
                   'Nature of Appointment (Regular/Fixed term/Temporary/Adhoc)', 'Nature of Duties', 
                   'Date of Joining', 'Date of Leaving', 'Salary with Grade', 'Reason of Leaving']
        
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for appt in appointments:
            row_cells = table.add_row().cells
            row_cells[0].text = appt.get('designation', '')
            row_cells[1].text = appt.get('employer', '')
            row_cells[2].text = appt.get('qualifications', '')
            row_cells[3].text = appt.get('nature', '')
            row_cells[4].text = appt.get('duties', '')
            row_cells[5].text = appt.get('joining_date', '')
            row_cells[6].text = appt.get('leaving_date', '')
            row_cells[7].text = appt.get('salary', '')
            row_cells[8].text = appt.get('reason', '')
    
    def add_current_posts_table(self, posts):
        """Add current institution posts table (Point 12)"""
        p = self.doc.add_paragraph()
        p.add_run('12. Posts Held after appointment at this Institution:').bold = True
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Designation', 'Department', 'Date of Joining', 'Grade Pay/ Pay Matrix Level']
        header_cells = table.rows[0].cells
        header_cells[0].text = headers[0]
        header_cells[1].text = headers[1]
        
        # Split Date of Joining into From/To sub-columns
        date_cell = header_cells[2]
        date_cell.text = headers[2]
        p_date = date_cell.paragraphs[0]
        p_date.add_run('\nFrom').bold = True
        p_date.add_run('\t\tTo').bold = True
        
        header_cells[3].text = headers[3]
        
        for paragraph in header_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        
        # Add data rows
        for post in posts:
            row_cells = table.add_row().cells
            row_cells[0].text = post.get('designation', '')
            row_cells[1].text = post.get('department', '')
            row_cells[2].text = f"{post.get('from_date', '')}\t\t{post.get('to_date', '')}"
            row_cells[3].text = post.get('grade_pay', '')
    
    def add_additional_fields(self, data):
        """Add fields 13-16"""
        # 13. Teaching experience
        p = self.doc.add_paragraph()
        p.add_run('13. Period of teaching experience:').bold = True
        p.add_run('\n\nP.G. Classes (In Years): ' + data.get('pg_experience', ''))
        p.add_run('\nU.G. Classes (In Years): ' + data.get('ug_experience', ''))
        
        # 14. Research experience
        p = self.doc.add_paragraph()
        p.add_run('14. Research Experience excluding Years Spent in M.Phil./Ph.D. (in Years)').bold = True
        p.add_run('\n' + data.get('research_experience', ''))
        
        # 15. Fields of specialization
        p = self.doc.add_paragraph()
        p.add_run('15. Fields of Specialization under the Subject / Discipline:').bold = True
        p.add_run('\n' + data.get('specialization', ''))
        
        # 16. Courses attended
        p = self.doc.add_paragraph()
        p.add_run('16. Human Resource Development Center Orientation / Refresher Course / FDP/ MOOC/ One-Two week Course attended so far:').bold = True
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Name of the Course', 'Place', 'Duration', 'Name of Organizer']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add course data if provided
        for course in data.get('courses', []):
            row_cells = table.add_row().cells
            row_cells[0].text = course.get('name', '')
            row_cells[1].text = course.get('place', '')
            row_cells[2].text = course.get('duration', '')
            row_cells[3].text = course.get('organizer', '')
        
        # Signature line
        p = self.doc.add_paragraph('\n\n')
        p.add_run('Name & Signature off Teacher').bold = False
    
    def add_part_b_header(self):
        """Add Part B header"""
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.add_run('PART B: ACADEMICPERFORMANCE INDICATORS (API):').bold = True
        
        intro = self.doc.add_paragraph()
        intro_text = """Based on the teacher's self-assessment, API score are proposed for (l ) teaching related activities; domain knowledge;

(2) Involvement in University / College student's related activities / research activities. The minimum API score required by teachers from this category is different for different levels of promotion. The self- assessment score should be based on objectively verifiable records. It shall be finalized by the Screening cum Evaluation /Selection Committee .University may detail the activities, incase institutional specificities require, and adjust the weight ages without changing the minimum total API scores required under this category

Tablet

Assessment Criteria and Methodology for University/College Teachers"""
        intro.add_run(intro_text)
    
    def add_teaching_table(self, teaching_data):
        """Add teaching performance table (Section 1)"""
        p = self.doc.add_paragraph()
        p.add_run('1.Teaching').bold = True
        
        # Create main table
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Main headers
        headers = ['Category', 'Name of Activity', 'Unit of Calculation', 'Self-Appraisal Grading', 
                   'Verified API Grading by Committee']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add sub-headers row
        row_cells = table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = ''
        row_cells[2].text = 'Actual Classes spent per Year\n\n% of Teaching'
        row_cells[3].text = """For Assistant Professor/ Associate Professor/ Professor

i)Good: 80% & Above

ii)Satisfactory : Below 80% but 70 % & Above

iii)Not satisfactory: Less than 70%"""
        row_cells[4].text = ''
        
        # Teaching activity row
        row_cells = table.add_row().cells
        row_cells[0].text = '1'
        row_cells[1].text = """Teaching: (Number of classes taught/total classes assigned) x100%(Classes

Taught includes sessions on tutorials,

lab and other

teaching related

activities)

(Teaching: Blackboard Teaching: ICT based Practical /Laboratory Tutorials /Assignments / Project, Field Work Group Discussion, Seminars Remedial! Teaching, Clarifying doubts within and outside the class hours Additional teaching to

Support counseling and mentoring)"""
        row_cells[2].text = teaching_data.get('actual_classes', '')
        row_cells[3].text = teaching_data.get('self_grading', '')
        row_cells[4].text = teaching_data.get('verified_grading', '')
        
        # Total row
        row_cells = table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = 'Total Actual hours spent'
        row_cells[1].merge(table.rows[-1].cells[2])
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    def add_activities_table(self, activities_data):
        """Add involvement in activities table (Section 2)"""
        p = self.doc.add_paragraph()
        p.add_run('2. Involvement in the University/College students related activities/research activities').bold = True
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Activities', 'Total days Spent per year', 'Self-Appraisal Grading For Assistant Professor/Associate Professor/Professor\n\ni)Good: 80% & Above\n\nii)Satisfactory : Below 80% but 70 % & Above\n\niii)Not satisfactory: Less than 70%', 
                   'Verified API Grading by Committee']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Column numbers row
        row_cells = table.add_row().cells
        row_cells[0].text = '(1)'
        row_cells[1].text = '(2)'
        row_cells[2].text = '(3)'
        row_cells[3].text = '(3)'
        
        # Activity rows
        activity_items = [
            ('(a) Administrative responsibilities such has Head, Chairperson /Dean/Director / IQAC Coordinator/different committees/Warden, etc.', 
             activities_data.get('admin_days', ''), activities_data.get('admin_grading', ''), ''),
            ('(b) Examination and evaluation duties assigned by the college/ university or attending the examination paper evaluation.\n\ni) Question Paper Setting\n\nii) Invigilation/Supervision\n\niii) Flying Squad\n\niv) CS/ACS/Custodian\n\nv) CAP Director l Assistant Director\n\nvi) Unfair Menace Committee\nvii)Grievance Committee\nviii)internal Assessment\n\nix) External Assessment\n\nx) Re-valuation\n\nxi) Result Preparation(College Level for Internal Assessment)\n\nxii) RRCIRAC Committee\n\nxiii) MPhil.,Ph.D. Thesis evaluation/any other)',
             activities_data.get('exam_days', ''), activities_data.get('exam_grading', ''), ''),
            ('(c)Student related co-curricular, extension and field based activities such as student clubs, career counselling, study visits, student seminars and other events, cultural, sports, NCC, NSS and community services.',
             activities_data.get('student_days', ''), activities_data.get('student_grading', ''), ''),
        ]
        
        for activity, days, grading, verified in activity_items:
            row_cells = table.add_row().cells
            row_cells[0].text = activity
            row_cells[1].text = days
            row_cells[2].text = grading
            row_cells[3].text = verified
        
        # Add remaining activities (d through g)
        remaining_activities = [
            'd) Institutional! Governance/ Participation in State/Central bodies/Committee on education, Research and National development etc.(Govt. Nominee/Nodal officer/Enquiry committee member/inspection committee member/state Govt. Workshop committee/Govt. CAS Committee/Subject expect/',
            '(d) Organizing seminars/ conferences/workshops, etc. and other college/university Activities.',
            '(e) Evidence of actively involved in guiding Ph.D. students\n\nNo. of Registered candidate:\n\nii)No. of Awarded Candidates:',
            '(f) Conducting Minor or Major Research Project sponsored by national or international agencies.\n\ni} Above 10 Lacs :\n\nii)Below 10lacs',
            '(g)At least one single or joint publication in peer-reviewed or UGC list of Journals.'
        ]
        
        for activity in remaining_activities:
            row_cells = table.add_row().cells
            row_cells[0].text = activity
            row_cells[1].text = activities_data.get(f'activity_{len(remaining_activities)}_days', '')
            row_cells[2].text = activities_data.get(f'activity_{len(remaining_activities)}_grading', '')
            row_cells[3].text = ''
        
        # Overall grading row
        row_cells = table.add_row().cells
        row_cells[0].text = 'Overall Grading:'
        row_cells[0].merge(row_cells[1])
        row_cells[2].text = """Good: Good in teaching and satisfactory or good in activity at S.No.2.Or Satisfactory: Satisfactory in teaching and good or satisfactory in activity At S.No.2.

Not Satisfactory: lf neither good nor satisfactory in overall grading."""
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    def add_note(self):
        """Add note at the end"""
        p = self.doc.add_paragraph()
        p.add_run('Note: ').bold = True
        note_text = """For the purpose of assessing the grading of Activity at Serial No. l and SerialNo.2, all such periods of duration which have been spent by the teacher on different kinds of paid leaves such as Maternity Leave, Child Care Leave, Study Leave, Medical Leave, Extraordinary Leave and Deputation shall be excluded from the grading assessment. The teacher shall be assessed for the remaining period of duration and the same shall be extrapolated forth entire period of assessment to arrive at the grading of the teacher. The teacher on such leaves or deputation as mentioned above shall not be put to any disadvantage for promotion under CAS due to his/her absence from his/her teaching responsibilities subject to the condition that such leave / deputation was undertaken with the prior approval of the competent authority following all procedures laid down in these regulations and as per the acts, statutes and ordinances of the parent institution"""
        p.add_run(note_text)
    
    def generate(self, data):
        """Generate complete PBAS form"""
        self.add_title()
        self.add_basic_info_section(data)
        self.add_part_a_header()
        self.add_general_info_fields(data)
        self.add_academic_qualifications_table(data.get('qualifications', []))
        self.add_research_degrees_table(data.get('research_degrees', []))
        self.add_prior_appointments_table(data.get('prior_appointments', []))
        self.add_current_posts_table(data.get('current_posts', []))
        self.add_additional_fields(data)
        self.add_part_b_header()
        self.add_teaching_table(data.get('teaching_data', {}))
        self.add_activities_table(data.get('activities_data', {}))
        self.add_note()
        
        return self.doc
    
    def save(self, filename='PBAS_Form_Output.docx'):
        """Save document"""
        self.doc.save(filename)


# Example usage with dummy data
if __name__ == '__main__':
    dummy_data = {
        'institute': 'ABC University',
        'department': 'Computer Science',
        'faculty': 'Engineering',
        'academic_year': '2024-2025',
        'name': 'DR. JOHN DOE',
        'current_designation': 'Assistant Professor, Level 10',
        'last_promotion': '01/07/2020',
        'cas_level': 'Level 11',
        'applied_designation': 'Assistant Professor (Senior Scale), Level 12',
        'eligibility_date': '01/07/2023',
        'address': '123 University Road, City - 400001',
        'mobile': '+91-9876543210',
        'email': 'john.doe@university.edu',
        'qualifications': [
            {'examination': 'S.S.C.', 'board': 'State Board', 'year': '2000', 'percentage': '85%', 'division': 'First', 'subject': 'All Subjects'},
            {'examination': 'H.S.C.', 'board': 'State Board', 'year': '2002', 'percentage': '82%', 'division': 'First', 'subject': 'Science'},
            {'examination': 'B.Sc.', 'board': 'XYZ University', 'year': '2005', 'percentage': '78%', 'division': 'First', 'subject': 'Computer Science'},
            {'examination': 'M.Sc.', 'board': 'XYZ University', 'year': '2007', 'percentage': '80%', 'division': 'First', 'subject': 'Computer Science'},
        ],
        'research_degrees': [
            {'title': 'Machine Learning Applications', 'date': '2012', 'university': 'ABC University'},
        ],
        'prior_appointments': [],
        'current_posts': [
            {'designation': 'Assistant Professor', 'department': 'Computer Science', 'from_date': '01/07/2013', 'to_date': 'Present', 'grade_pay': 'Level 10'},
        ],
        'pg_experience': '11 years',
        'ug_experience': '11 years',
        'research_experience': '8 years',
        'specialization': 'Machine Learning, Artificial Intelligence, Data Mining',
        'courses': [
            {'name': 'Refresher Course in AI', 'place': 'Mumbai', 'duration': '2 weeks', 'organizer': 'UGC-HRDC'},
            {'name': 'FDP on Data Science', 'place': 'Pune', 'duration': '1 week', 'organizer': 'AICTE'},
        ],
        'teaching_data': {
            'actual_classes': '240 hours\n\n95%',
            'self_grading': 'Good',
            'verified_grading': '',
        },
        'activities_data': {
            'admin_days': '30',
            'admin_grading': 'Good',
            'exam_days': '45',
            'exam_grading': 'Good',
            'student_days': '25',
            'student_grading': 'Satisfactory',
        }
    }
    
    generator = PBASFormGenerator()
    generator.generate(dummy_data)
    generator.save('PBAS_Form_Generated.docx')
    print("PBAS Form generated successfully!")